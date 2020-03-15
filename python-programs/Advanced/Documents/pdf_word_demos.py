import PyPDF2
import docx
import os

os.chdir(r'D:\Knowledge\Python\git\python-programs\Advanced\Documents')
pdf_doc = open(r'meetingminutes1.pdf', 'rb')

pdf_reader = PyPDF2.PdfFileReader(pdf_doc)

pdf_string = ''
for page_index in range(pdf_reader.getNumPages()):
    page = pdf_reader.getPage(page_index)
    pdf_string += page.extractText()
    
print(pdf_string)
pdf_doc.close()

pdf_file1 = open(r'meetingminutes1.pdf', 'rb')
pdf_file2 = open(r'meetingminutes2.pdf', 'rb')
pdf_reader1 = PyPDF2.PdfFileReader(pdf_file1)
pdf_reader2 = PyPDF2.PdfFileReader(pdf_file2)

writer = PyPDF2.PdfFileWriter()
for index in range(pdf_reader1.numPages):
    page = pdf_reader1.getPage(index)
    writer.addPage(page)

for index in range(pdf_reader2.numPages):
    page = pdf_reader2.getPage(index)
    writer.addPage(page)

output = open('merged_pdf.pdf', 'wb')
writer.write(output)
output.close()
pdf_file1.close()
pdf_file2.close()

d = docx.Document()
d.add_paragraph(pdf_string)
p = d.paragraphs[0]
p.runs[0].bold = True
d.save('my_doc.docx')
